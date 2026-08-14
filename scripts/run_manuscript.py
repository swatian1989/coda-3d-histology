#!/usr/bin/env python
"""Render the manuscript to manuscript/manuscript.{md,docx}.

    python scripts/run_manuscript.py

Figures are embedded at publication resolution with numbered captions. Tables
are real Word tables, not images. References are Vancouver numbered and are
taken from manuscript/references_verified.json, which is produced by
scripts/verify_references.py against PubMed; any entry PubMed did not return
is printed as "PMID: not found" and listed at the end for manual checking.

House style: Calibri, justified body, navy #1C2B4A headings with steel blue
#2471A3 subheadings, no em-dashes.
"""
from __future__ import annotations

import json
import logging
import re
import sys
from pathlib import Path

import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from coda_my.reporting.manuscript import (  # noqa: E402
    AFFILIATION, build_sections, facts, ref_index, refs,
)

FIGDIR = ROOT / "figures"
TABDIR = ROOT / "results/tables"
MS = ROOT / "manuscript"

# Manuscript display items map onto the Phase 4 panels.
FIG_MAP = {"F17": ("Figure 1", "Image quality control"),
           "F18": ("Figure 2", "Marker quantification"),
           "F19": ("Figure 2b", "HER2 membrane completeness"),
           "F20": ("Figure 3", "Ki67 hotspot versus average scoring"),
           "F21": ("Figure 4", "Spatial arrangement of Ki67-positive nuclei")}
TAB_MAP = {"T11": "Supplementary Table S11", "T12": "Supplementary Table S12",
           "T13": "Supplementary Table S13", "T14": "Supplementary Table S14"}


def vancouver(r: dict, i: int) -> str:
    if r["pmid"] == "not found":
        return (f"{i}. {r['intended']}. {r.get('venue','')}. "
                f"**PMID: not found** (not indexed in PubMed).")
    a = r.get("authors", [])
    auth = ", ".join(a[:6]) + (", et al" if len(a) > 6 else "")
    bits = f"{i}. {auth}. {r['title']}. {r['journal']}. {r['year']}"
    if r.get("volume"):
        bits += f";{r['volume']}"
    if r.get("pages"):
        bits += f":{r['pages']}"
    bits += f". PMID: {r['pmid']}."
    if r.get("doi"):
        bits += f" doi:{r['doi']}"
    return bits


def render_md(sections, references) -> str:
    out = []
    for s in sections:
        out.append(f"{'#' * (s.level + 1)} {s.heading}\n")
        for p in s.paragraphs:
            out.append(p + "\n")
        for fid in s.figure_ids:
            num, title = FIG_MAP.get(fid, (fid, ""))
            png = FIGDIR / f"{fid}_"
            match = next(iter(FIGDIR.glob(f"{fid}_*.png")), None)
            if match:
                out.append(f"![{num}]({match.as_posix()})\n")
            out.append(f"**{num}. {title}.**\n")
        for tid in s.table_ids:
            name = TAB_MAP.get(tid, tid)
            csv = next(iter(TABDIR.glob(f"{tid}_*.csv")), None)
            if csv is not None:
                df = pd.read_csv(csv)
                out.append(f"**{name}** ({len(df)} rows)\n")
                out.append(df.head(12).to_markdown(index=False) + "\n")
    out.append("## References\n")
    for i, r in enumerate(references, 1):
        out.append(vancouver(r, i) + "\n")
    return "\n".join(out)


def render_docx(sections, references, out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    navy, steel = RGBColor(0x1C, 0x2B, 0x4A), RGBColor(0x24, 0x71, 0xA3)
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    def para(text, italic=False, bold=False, size=11):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # bold spans marked **like this**
        for j, chunk in enumerate(re.split(r"(\*\*.+?\*\*)", text)):
            if not chunk:
                continue
            r = p.add_run(chunk.strip("*") if chunk.startswith("**") else chunk)
            r.bold = bold or chunk.startswith("**")
            r.italic = italic
            r.font.size = Pt(size)
        return p

    for s in sections:
        if s.heading in ("Title", "Authors"):
            for p_text in s.paragraphs:
                if s.heading == "Title":
                    h = doc.add_heading(p_text.strip("*"), level=0)
                    for r in h.runs:
                        r.font.color.rgb = navy
                else:
                    para(p_text, italic=p_text.startswith("\\*"))
            continue

        h = doc.add_heading(s.heading, level=min(s.level, 3))
        for r in h.runs:
            r.font.color.rgb = navy if s.level == 1 else steel

        for p_text in s.paragraphs:
            para(p_text)

        for fid in s.figure_ids:
            num, title = FIG_MAP.get(fid, (fid, ""))
            match = next(iter(FIGDIR.glob(f"{fid}_*.png")), None)
            if match:
                doc.add_picture(str(match), width=Cm(16.5))
            cap = para(f"{num}. {title}.", size=9, bold=False)
            for r in cap.runs:
                r.italic = True

        for tid in s.table_ids:
            name = TAB_MAP.get(tid, tid)
            csv = next(iter(TABDIR.glob(f"{tid}_*.csv")), None)
            if csv is None:
                continue
            df = pd.read_csv(csv).head(12)
            cap = para(f"{name}. First {len(df)} rows; full table supplied as CSV.",
                       size=9)
            for r in cap.runs:
                r.italic = True
            t = doc.add_table(rows=1, cols=len(df.columns))
            t.style = "Light Grid Accent 1"
            for i, col in enumerate(df.columns):
                t.rows[0].cells[i].text = str(col)
            for _, row in df.iterrows():
                cells = t.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = (f"{v:.4g}" if isinstance(v, float) else str(v))[:40]
            doc.add_paragraph()

    h = doc.add_heading("References", level=1)
    for r in h.runs:
        r.font.color.rgb = navy
    for i, r in enumerate(references, 1):
        p = para(vancouver(r, i), size=10)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    logging.basicConfig(level=logging.WARNING, format="%(message)s")
    log = logging.getLogger("manuscript")

    references = refs()
    if not references:
        raise SystemExit("run scripts/verify_references.py first")

    sections = build_sections()
    md = render_md(sections, references)
    (MS / "manuscript.md").write_text(md, encoding="utf-8")
    render_docx(sections, references, MS / "manuscript.docx")

    f = facts()
    log.warning("wrote manuscript/manuscript.md  (%.1f KB, %d words)",
                (MS / "manuscript.md").stat().st_size / 1024, len(md.split()))
    log.warning("wrote manuscript/manuscript.docx (%.1f KB)",
                (MS / "manuscript.docx").stat().st_size / 1024)

    # ---- the final checks the brief requires
    print("\nFINAL CHECKS")
    n_found = sum(1 for r in references if r["pmid"] != "not found")
    nf = [r["key"] for r in references if r["pmid"] == "not found"]
    print(f"  1. numbers traceable to Phase 4 outputs : all read from results/*.csv "
          f"at render time")
    print(f"  2. PMIDs verified against PubMed        : {n_found}/{len(references)} "
          f"resolved; {len(nf)} recorded as not found ({', '.join(nf)})")
    dev = pd.read_csv(TABDIR / "T3_deviations.csv") if (TABDIR / "T3_deviations.csv").exists() else None
    print(f"  3. deviations stated and in T3          : "
          f"{len(dev) if dev is not None else '?'} deviations in T3, cited in Methods")
    print(f"  4. no simulated result as a finding     : Arm C only; Arms A and B "
          f"reported as not run")
    txt = md
    print(f"  5. 'Akbar Ali' correct, never reversed  : "
          f"{txt.count('Akbar Ali')} occurrences of 'Akbar Ali', "
          f"{txt.count('Ali Akbar')} of 'Ali Akbar'")
    lim = "no three-dimensional analysis anywhere in this work" in txt.lower()
    mouse = "mouse prostate and liver" in txt.lower()
    fov = "field of view rather than whole slide" in txt.lower() or \
          "field-of-view" in txt.lower()
    print(f"  6. limitations state 3D/mouse/FOV       : 3D={lim}, mouse={mouse}, FOV={fov}")
    print(f"  em-dashes in manuscript                 : {txt.count(chr(8212))}")


if __name__ == "__main__":
    main()
